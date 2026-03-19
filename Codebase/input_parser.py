"""
input_parser.py - Handles reading resume and job description inputs
from multiple formats: PDF, DOCX, TXT, and manual terminal input.

Supported input methods:
- PDF files (using pdfplumber for robust text extraction)
- DOCX files (using python-docx for paragraph and table extraction)
- Plain text files (with multi-encoding support)
- Manual free-form text input from terminal
- Guided section-by-section input from terminal
"""

import os
import sys


class InputParser:
    """
    Parses input documents (resume or job description) from various file formats
    into plain text strings suitable for LLM processing.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".text"}

    def read_file(self, file_path):
        """
        Detect file type by extension and dispatch to the appropriate parser.

        Args:
            file_path (str): Path to the input file

        Returns:
            str: Extracted plain text content

        Raises:
            FileNotFoundError: If the file does not exist at the given path
            ValueError: If the file extension is not supported
        """
        # Validate file exists
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect extension and dispatch
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported formats: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        if ext == ".pdf":
            return self.read_pdf(file_path)
        elif ext in (".docx", ".doc"):  # .doc is best-effort; python-docx may fail on legacy OLE format
            return self.read_docx(file_path)
        elif ext in (".txt", ".text"):
            return self.read_text_file(file_path)

    def read_pdf(self, file_path):
        """
        Extract text from a PDF file using pdfplumber.
        Handles multi-page documents and preserves paragraph structure.

        Args:
            file_path (str): Path to the PDF file

        Returns:
            str: Full text content concatenated from all pages
        """
        try:
            import pdfplumber
        except ImportError:
            print("ERROR: 'pdfplumber' package is not installed.")
            print("Please install it with: pip install pdfplumber")
            sys.exit(1)

        try:
            with pdfplumber.open(file_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
        except Exception as e:
            # Corrupt, password-protected, or otherwise unreadable PDF
            raise ValueError(
                f"Could not read PDF '{file_path}': {e}"
            ) from e

        # Double newline between pages preserves visual section breaks for LLM extraction
        result = "\n\n".join(pages_text).strip()
        if not result:
            print(f"  Warning: No text could be extracted from '{file_path}'.")
            print("  The PDF may be image-based or encrypted.")
        return result

    def read_docx(self, file_path):
        """
        Extract text from a DOCX file using python-docx.
        Reads both paragraphs and tables to capture all content,
        as many resume templates use invisible tables for layout.

        Args:
            file_path (str): Path to the DOCX file

        Returns:
            str: Full text content from paragraphs and tables
        """
        try:
            from docx import Document
        except ImportError:
            print("ERROR: 'python-docx' package is not installed.")
            print("Please install it with: pip install python-docx")
            sys.exit(1)

        try:
            doc = Document(file_path)
        except Exception as e:
            # Corrupt, password-protected, or otherwise unreadable DOCX
            raise ValueError(
                f"Could not read DOCX '{file_path}': {e}"
            ) from e
        all_text = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())

        # Extract table content (resumes often use tables for multi-column layout)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    # Deduplicate cells (merged cells repeat content)
                    seen = []
                    for cell_text in row_data:
                        if cell_text not in seen:
                            seen.append(cell_text)
                    all_text.append(" | ".join(seen))  # pipe-delimited flat layout for LLM parsing

        result = "\n".join(all_text)
        if not result:
            print(f"  Warning: No text could be extracted from '{file_path}'.")
        return result

    def read_text_file(self, file_path):
        """
        Read plain text file content with multiple encoding fallbacks.
        Tries UTF-8 first, then falls back to common encodings.

        Args:
            file_path (str): Path to the text file

        Returns:
            str: File content as string

        Raises:
            ValueError: If the file cannot be read with any supported encoding
        """
        # Try progressively more permissive encodings; latin-1 accepts all bytes as last resort
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                continue

        raise ValueError(
            f"Unable to read '{file_path}' with any supported encoding: "
            f"{', '.join(encodings)}"
        )

    def read_manual_input(self):
        """
        Collect free-form multi-line text input from the terminal.
        User types their resume content and presses Enter twice on an empty line
        to finish input.

        Returns:
            str: User-entered text content
        """
        print("\n  Enter your resume information below.")
        print("  (Press Enter twice on an empty line to finish)\n")

        lines = []
        empty_line_count = 0

        while True:
            try:
                line = input()
                if line.strip() == "":
                    empty_line_count += 1
                    if empty_line_count >= 2:
                        break
                    lines.append("")  # Preserve single blank lines
                else:
                    empty_line_count = 0
                    lines.append(line)
            except EOFError:
                break

        result = "\n".join(lines).strip()
        if not result:
            print("  Warning: No input was provided.")
        return result

    def read_manual_structured(self):
        """
        Guided manual input with adaptive branching based on profile type.
        After collecting contact info and education, asks the user whether
        they are a fresher or experienced professional and adjusts the
        remaining questions accordingly.

        Fresher path: skips work experience, asks about internships (light),
        emphasizes projects and skills.

        Experienced path: asks for years of experience, full work experience
        loop, optional projects.

        Returns:
            str: Concatenated text from all prompted sections
        """
        print("\n  === Guided Resume Input ===")
        print("  (Press Enter to skip any field)\n")

        # ── 1. Contact Info (always) ───────────────────────────────────────
        name = input("  Full Name: ").strip()
        email = input("  Email: ").strip()
        phone = input("  Phone: ").strip()
        location = input("  City, Country: ").strip()
        linkedin = input("  LinkedIn URL (optional): ").strip()
        portfolio = input("  Portfolio/GitHub URL (optional): ").strip()

        # ── 2. Education (always) ──────────────────────────────────────────
        print("\n  --- Education (press Enter on Degree to stop adding) ---")
        education_entries = []
        while True:
            degree = input("  Degree (e.g., MBBS, B.E. Mechanical, MBA): ").strip()
            if not degree:
                break
            institution = input("  Institution: ").strip()
            year = input("  Year of completion: ").strip()
            gpa = input("  Score - CGPA, GPA, or Percentage (optional, e.g., 8.75/10, 3.8/4.0, 72%): ").strip()
            entry = f"{degree} - {institution} ({year})"
            if gpa:
                # >10 = percentage (GPA/CGPA scales top out at 4.0 and 10.0)
                if "%" in gpa or (gpa.replace(".", "").isdigit() and float(gpa) > 10):
                    entry += f" | Score: {gpa}" if "%" in gpa else f" | Score: {gpa}%"
                elif "/" in gpa:
                    entry += f" | CGPA: {gpa}"
                else:
                    entry += f" | CGPA: {gpa}"
            education_entries.append(entry)
            print()

        # ── 3. Branch Point ────────────────────────────────────────────────
        print("\n  What best describes you?")
        print("  [1] Fresher / Recent Graduate (no full-time work experience)")
        print("  [2] Experienced Professional")

        profile_type = ""
        while profile_type not in ("1", "2"):
            profile_type = input("\n  Enter choice (1-2): ").strip()
            if profile_type not in ("1", "2"):
                print("  Invalid choice. Please enter 1 or 2.")

        is_fresher = profile_type == "1"

        experience_entries = []
        years_exp = ""

        if is_fresher:
            # ── Fresher: Internships (light format) ────────────────────────
            print("\n  --- Internships / Volunteer Work (press Enter to skip) ---")
            print("  (Lighter format — just the basics)")
            while True:
                role = input("  Role/Title: ").strip()
                if not role:
                    break
                org = input("  Organization: ").strip()
                what = input("  What did you do? (1-2 sentences): ").strip()
                entry = f"{role} at {org}"
                if what:
                    entry += f"\n  - {what}"
                experience_entries.append(entry)
                print()
        else:
            # ── Experienced: years + full work experience ──────────────────
            years_exp = input("\n  How many years of professional experience do you have? ").strip()

            print("\n  --- Work Experience (press Enter on Job Title to stop adding) ---")
            while True:
                title = input("  Job Title: ").strip()
                if not title:
                    break
                company = input("  Company: ").strip()
                dates = input("  Dates (e.g., Jan 2020 - Dec 2023): ").strip()
                print("  Key responsibilities/achievements (one per line, blank to finish):")
                responsibilities = []
                while True:
                    resp = input("    - ").strip()
                    if not resp:
                        break
                    responsibilities.append(resp)

                entry = f"{title} at {company} ({dates})"
                if responsibilities:
                    entry += "\n" + "\n".join(f"  - {r}" for r in responsibilities)
                experience_entries.append(entry)
                print()

        # ── 4. Skills (always) ─────────────────────────────────────────────
        print("\n  --- Skills ---")
        tech_skills = input("  Technical skills (comma-separated): ").strip()
        soft_skills = input("  Soft skills (comma-separated): ").strip()
        tools = input("  Tools/Software (comma-separated): ").strip()
        languages = input("  Programming/Human languages (comma-separated): ").strip()

        # Re-prompt if all skill fields are empty
        if not any([tech_skills, soft_skills, tools, languages]):
            print("\n  Skills help your CV stand out. Please list at least a few.")
            tech_skills = input("  Technical skills (comma-separated): ").strip()

        # ── 5. Projects ───────────────────────────────────────────────────
        if is_fresher:
            print("\n  --- Projects (important for your profile!) ---")
            print("  Projects are key for freshers. Please describe at least one.")
        else:
            print("\n  --- Projects (optional, press Enter to skip) ---")

        project_entries = []
        while True:
            proj_name = input("  Project Name: ").strip()
            if not proj_name:
                if is_fresher and not project_entries:
                    print("  We strongly recommend adding at least one project.")
                    proj_name = input("  Project Name: ").strip()
                    if not proj_name:
                        break
                else:
                    break
            proj_desc = input("  Brief Description: ").strip()
            proj_tech = input("  Technologies used: ").strip()
            entry = f"{proj_name}: {proj_desc} (Technologies: {proj_tech})"
            project_entries.append(entry)
            print()

        # ── 6. Certifications (optional) ──────────────────────────────────
        print("\n  --- Certifications (press Enter to skip) ---")
        certifications = []
        while True:
            cert = input("  Certification (e.g., PMP, CPA, Six Sigma): ").strip()
            if not cert:
                break
            certifications.append(cert)

        # ── 7. Achievements (optional) ────────────────────────────────────
        if is_fresher:
            print("\n  --- Achievements (academic awards, competitions, hackathons) ---")
        else:
            print("\n  --- Achievements (press Enter to skip) ---")
        achievements = []
        while True:
            ach = input("  Achievement: ").strip()
            if not ach:
                break
            achievements.append(ach)

        # ── Combine into text block ───────────────────────────────────────
        combined_parts = [f"Name: {name}"]

        contact_line_parts = []
        if email:
            contact_line_parts.append(f"Email: {email}")
        if phone:
            contact_line_parts.append(f"Phone: {phone}")
        if location:
            contact_line_parts.append(f"Location: {location}")
        if linkedin:
            contact_line_parts.append(f"LinkedIn: {linkedin}")
        if portfolio:
            contact_line_parts.append(f"Portfolio: {portfolio}")
        if contact_line_parts:
            combined_parts.append(" | ".join(contact_line_parts))

        if is_fresher:
            combined_parts.append("\nPROFILE TYPE: Fresher / Recent Graduate")
            combined_parts.append("YEARS OF EXPERIENCE: 0")
        else:
            combined_parts.append("\nPROFILE TYPE: Experienced Professional")
            if years_exp:
                combined_parts.append(f"YEARS OF EXPERIENCE: {years_exp}")

        if education_entries:
            combined_parts.append("\nEDUCATION:")
            combined_parts.extend(education_entries)

        if experience_entries:
            if is_fresher:
                combined_parts.append("\nINTERNSHIPS / VOLUNTEER WORK:")
            else:
                combined_parts.append("\nWORK EXPERIENCE:")
            combined_parts.extend(experience_entries)

        skill_parts = []
        if tech_skills:
            skill_parts.append(f"Technical: {tech_skills}")
        if soft_skills:
            skill_parts.append(f"Soft Skills: {soft_skills}")
        if tools:
            skill_parts.append(f"Tools: {tools}")
        if languages:
            skill_parts.append(f"Languages: {languages}")
        if skill_parts:
            combined_parts.append("\nSKILLS:")
            combined_parts.extend(skill_parts)

        if project_entries:
            combined_parts.append("\nPROJECTS:")
            combined_parts.extend(project_entries)

        if certifications:
            combined_parts.append("\nCERTIFICATIONS:")
            combined_parts.extend(certifications)

        if achievements:
            combined_parts.append("\nACHIEVEMENTS:")
            combined_parts.extend(achievements)

        return "\n".join(combined_parts)

    def read_job_description_structured(self):
        """
        Guided section-by-section input for a job description.
        Collects key JD fields to produce a well-structured text block
        that the extraction model can parse reliably.

        Returns:
            str: Concatenated text from all prompted sections
        """
        print("\n  === Guided Job Description Input ===")
        print("  (Press Enter to skip any field)\n")

        # ── Basic Info ────────────────────────────────────────────────
        job_title = input("  Job Title: ").strip()
        company = input("  Company Name: ").strip()
        location = input("  Location (e.g., Remote, New York, Hybrid): ").strip()
        employment_type = input("  Employment Type (Full-time / Part-time / Contract): ").strip()

        # ── Experience & Education ────────────────────────────────────
        experience = input("  Required Years of Experience (e.g., 3+, 5-7): ").strip()
        education = input("  Education Requirement (e.g., Bachelor's degree, MBA): ").strip()

        # ── Skills ────────────────────────────────────────────────────
        print("\n  --- Skills ---")
        required_skills = input("  Required Skills (comma-separated): ").strip()
        preferred_skills = input("  Preferred / Nice-to-Have Skills (comma-separated): ").strip()
        tools = input("  Tools / Technologies (comma-separated): ").strip()

        # ── Responsibilities ──────────────────────────────────────────
        print("\n  --- Responsibilities (one per line, blank to finish) ---")
        responsibilities = []
        while True:
            resp = input("    - ").strip()
            if not resp:
                break
            responsibilities.append(resp)

        # ── Additional Info ───────────────────────────────────────────
        salary = input("\n  Salary Range (optional): ").strip()
        other = input("  Any Other Details (optional): ").strip()

        # ── Combine into text block ──────────────────────────────────
        parts = []

        if job_title:
            parts.append(f"Job Title: {job_title}")
        if company:
            parts.append(f"Company: {company}")
        if location:
            parts.append(f"Location: {location}")
        if employment_type:
            parts.append(f"Employment Type: {employment_type}")

        if experience:
            parts.append(f"\nRequired Experience: {experience}")
        if education:
            parts.append(f"Education Requirement: {education}")

        skill_parts = []
        if required_skills:
            skill_parts.append(f"Required Skills: {required_skills}")
        if preferred_skills:
            skill_parts.append(f"Preferred Skills: {preferred_skills}")
        if tools:
            skill_parts.append(f"Tools / Technologies: {tools}")
        if skill_parts:
            parts.append("\nSKILLS:")
            parts.extend(skill_parts)

        if responsibilities:
            parts.append("\nRESPONSIBILITIES:")
            parts.extend(f"- {r}" for r in responsibilities)

        if salary:
            parts.append(f"\nSalary Range: {salary}")
        if other:
            parts.append(f"\nAdditional Details: {other}")

        result = "\n".join(parts)
        if not result.strip():
            print("  Warning: No job description data was provided.")
        return result
