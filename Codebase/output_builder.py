"""
output_builder.py - Generates final CV documents in DOCX and PDF formats.
Uses python-docx for DOCX creation and fpdf2 for PDF generation.

Applies professional, ATS-friendly formatting with:
- Clean, readable fonts (Calibri for DOCX, Helvetica for PDF)
- Consistent section headings and bullet points
- Narrow margins for space efficiency
- Standard formatting that ATS systems can parse
"""

import os
import json
import sys


class OutputBuilder:
    """
    Builds professionally formatted CV documents from structured content.
    Supports DOCX and PDF output with ATS-friendly formatting.
    """

    # Placeholder values that should never appear in output
    _PLACEHOLDERS = {"fill_in", "fill in", "n/a", "na", "none", "nil", "-", ".", "none."}

    def __init__(self, output_dir="."):
        """
        Initialize the output builder.

        Args:
            output_dir (str): Directory to save generated files (default: current dir)
        """
        self.output_dir = output_dir

    @classmethod
    def _is_placeholder(cls, value):
        """
        Check whether a string value is a known placeholder that should
        not be rendered in the output document.

        Args:
            value: The value to check

        Returns:
            bool: True if the value is a placeholder
        """
        if not isinstance(value, str):
            return False
        return value.strip().lower() in cls._PLACEHOLDERS or not value.strip()

    @staticmethod
    def _parse_bold_segments(text):
        """
        Parse markdown **bold** markers into a list of (text, is_bold) segments.
        E.g. "Used **Python** and **React**" →
             [("Used ", False), ("Python", True), (" and ", False), ("React", True)]

        Args:
            text (str): Text potentially containing **bold** markers

        Returns:
            list: List of (text, is_bold) tuples
        """
        import re
        segments = []
        last_end = 0
        for match in re.finditer(r'\*\*(.+?)\*\*', text):
            # Add text before the bold marker
            if match.start() > last_end:
                segments.append((text[last_end:match.start()], False))
            # Add the bold text
            segments.append((match.group(1), True))
            last_end = match.end()
        # Add remaining text after the last bold marker
        if last_end < len(text):
            segments.append((text[last_end:], False))
        # If no bold markers found, return the whole text as non-bold
        if not segments:
            segments.append((text, False))
        return segments

    @staticmethod
    def _is_experience_title_line(line):
        """
        Detect whether a line looks like an experience header, e.g.:
          "Digital Marketing Manager | StyleHub E-commerce | 2021-03 - Present"
          "Marketing Manager | Company | March 2021 -"
        Pattern: contains at least one pipe separator AND contains a date-like
        token (year, "Present", month-year, or trailing dash after a year).

        Args:
            line (str): A single line of text

        Returns:
            bool: True if the line matches an experience title pattern
        """
        import re
        # Must have at least one pipe separator
        if "|" not in line:
            return False
        stripped = line.strip()
        # Ends with a year (optionally with -MM suffix) or "Present"
        if re.search(r'(\d{4}(-\d{2})?|[Pp]resent)\s*$', stripped):
            return True
        # "START - END" pattern (e.g. "2021 - 2023" or "Jan 2021 - Present")
        if re.search(r'\d{4}.*-\s*(Present|\d{4})', stripped):
            return True
        # Trailing dash after a year (e.g. "March 2021 -" meaning ongoing)
        if re.search(r'\d{4}\s*-\s*$', stripped):
            return True
        # Contains a date anywhere and has pipes (broad fallback)
        if re.search(r'\b(19|20)\d{2}\b', stripped) and stripped.count("|") >= 2:
            return True
        return False

    def build_docx(self, cv_content, contact_info, filename="generated_cv"):
        """
        Create a formatted DOCX file from CV content.
        Uses ATS-friendly fonts, clean headings, and bullet points.

        Args:
            cv_content (dict): Generated CV content organized by section
            contact_info (dict): Name, email, phone, location, linkedin, portfolio
            filename (str): Output filename without extension

        Returns:
            str: Full path to the generated DOCX file
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("ERROR: 'python-docx' package is not installed.")
            print("Please install it with: pip install python-docx")
            sys.exit(1)

        doc = Document()

        # Set narrow margins (0.7 inch on all sides)
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # --- Header: Candidate Name ---
        name = contact_info.get("name", "Candidate Name")
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(name.upper())
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.name = "Calibri"
        name_para.space_after = Pt(2)

        # --- Contact Line ---
        contact_parts = []
        for field in ["email", "phone", "location", "linkedin"]:
            val = contact_info.get(field, "")
            if val and not self._is_placeholder(val):
                contact_parts.append(val)
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(9)
            contact_run.font.name = "Calibri"
            contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            contact_para.space_after = Pt(4)

        # --- Horizontal Rule ---
        self._add_horizontal_rule(doc)

        # --- CV Sections ---
        section_order = [
            "professional_summary", "experience", "education",
            "skills", "projects", "certifications", "achievements"
        ]

        for section_name in section_order:
            if section_name not in cv_content:
                continue
            content = cv_content[section_name]

            # Skip empty sections
            if not content or content == "" or content == [] or content == {}:
                continue

            # Section heading
            heading_text = section_name.replace("_", " ").upper()
            self._add_section_heading_docx(doc, heading_text)

            # Format content based on type
            if isinstance(content, str):
                self._add_text_paragraph(doc, content, section_name=section_name)
            elif isinstance(content, list):
                self._add_list_content(doc, content, section_name)
            elif isinstance(content, dict):
                self._add_dict_content(doc, content)

        # Save the document
        output_path = os.path.join(self.output_dir, f"{filename}.docx")
        doc.save(output_path)
        return output_path

    def build_pdf(self, cv_content, contact_info, filename="generated_cv"):
        """
        Create a formatted PDF file from CV content using fpdf2.
        Mirrors the DOCX layout for consistency.

        Args:
            cv_content (dict): Generated CV content organized by section
            contact_info (dict): Name, email, phone, location, linkedin, portfolio
            filename (str): Output filename without extension

        Returns:
            str: Full path to the generated PDF file
        """
        try:
            from fpdf import FPDF
        except ImportError:
            print("ERROR: 'fpdf2' package is not installed.")
            print("Please install it with: pip install fpdf2")
            sys.exit(1)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(15, 15, 15)

        # --- Header: Candidate Name ---
        name = contact_info.get("name", "Candidate Name")
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, name.upper(), align="C", new_x="LMARGIN", new_y="NEXT")

        # --- Contact Line ---
        contact_parts = []
        for field in ["email", "phone", "location", "linkedin"]:
            val = contact_info.get(field, "")
            if val and not self._is_placeholder(val):
                contact_parts.append(val)
        if contact_parts:
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, self._sanitize_pdf_text(" | ".join(contact_parts)),
                     align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

        # --- Horizontal Rule ---
        pdf.ln(3)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

        # --- CV Sections ---
        section_order = [
            "professional_summary", "experience", "education",
            "skills", "projects", "certifications", "achievements"
        ]

        for section_name in section_order:
            if section_name not in cv_content:
                continue
            content = cv_content[section_name]

            # Skip empty sections
            if not content or content == "" or content == [] or content == {}:
                continue
            # Skip sections where all text lines are just "NONE" placeholders
            if isinstance(content, str):
                real_lines = [l.strip() for l in content.split("\n")
                              if l.strip() and not self._is_placeholder(l.strip())]
                if not real_lines:
                    continue

            # Section heading
            heading_text = section_name.replace("_", " ").upper()
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(31, 73, 125)  # Dark blue
            pdf.cell(0, 8, heading_text, new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

            # Thin line under heading
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(2)

            # Format content
            pdf.set_font("Helvetica", "", 10)
            if isinstance(content, str):
                self._add_text_pdf(pdf, content, section_name=section_name)
            elif isinstance(content, list):
                self._add_list_pdf(pdf, content, section_name)
            elif isinstance(content, dict):
                self._add_dict_pdf(pdf, content)

            pdf.ln(3)

        # Save the document
        output_path = os.path.join(self.output_dir, f"{filename}.pdf")
        pdf.output(output_path)
        return output_path

    # =========================================================================
    # DOCX Helper Methods
    # =========================================================================

    def _add_section_heading_docx(self, doc, heading_text):
        """
        Add a formatted section heading to the DOCX document.

        Args:
            doc: python-docx Document object
            heading_text (str): Section title text (will be displayed uppercase)
        """
        from docx.shared import Pt, RGBColor

        heading_para = doc.add_paragraph()
        heading_para.space_before = Pt(8)
        heading_para.space_after = Pt(2)
        heading_run = heading_para.add_run(heading_text)
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.name = "Calibri"
        heading_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue

    def _add_horizontal_rule(self, doc):
        """
        Add a horizontal line to the DOCX document.

        Args:
            doc: python-docx Document object
        """
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        para = doc.add_paragraph()
        para.space_before = Pt(0)
        para.space_after = Pt(4)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '999999')
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _sanitize_docx_text(text):
        """
        Clean LLM-generated text for DOCX output. Strips markdown italic
        markers (*text*, __text__) but KEEPS **bold** markers for rich
        text rendering via _add_rich_text_docx.

        Args:
            text (str): Raw text that may contain markdown markers

        Returns:
            str: Cleaned text with only **bold** markers preserved
        """
        import re
        if not text:
            return text
        # Keep **bold** markers, strip single *italic* and __underline__
        text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        return text

    def _add_rich_text_docx(self, doc, text, style=None):
        """
        Add a paragraph with mixed bold/regular text to the DOCX document.
        Parses **bold** markers and renders them as actual bold runs.

        Args:
            doc: python-docx Document object
            text (str): Text with optional **bold** markers
            style (str or None): Paragraph style (e.g. "List Bullet")

        Returns:
            Paragraph: The created paragraph object
        """
        from docx.shared import Pt

        segments = self._parse_bold_segments(self._sanitize_docx_text(text))

        if style:
            para = doc.add_paragraph(style=style)
        else:
            para = doc.add_paragraph()

        for segment_text, is_bold in segments:
            if not segment_text:
                continue
            run = para.add_run(segment_text)
            run.bold = is_bold
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"

        return para

    def _add_text_paragraph(self, doc, text, section_name=""):
        """
        Add a text paragraph to the DOCX document with proper formatting.
        Handles multi-line text by splitting on newlines.
        Renders **bold** markers as actual bold text for keywords.
        Detects experience title lines and renders them fully bold.

        Args:
            doc: python-docx Document object
            text (str): Paragraph text content
            section_name (str): Section name for context-aware formatting
        """
        from docx.shared import Pt
        import re

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Skip placeholder lines (NONE, N/A, etc.)
            if self._is_placeholder(line):
                continue
            # Skip markdown horizontal rules (---, ***, ___)
            if set(line) <= {'-', '*', '_'} and len(line) >= 3:
                continue
            # Detect experience title lines (e.g. "Title | Company | Date")
            if self._is_experience_title_line(line):
                para = doc.add_paragraph()
                para.space_before = Pt(4)
                para.space_after = Pt(1)
                run = para.add_run(self._sanitize_docx_text(line))
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"
                continue
            # Detect bullet points (-, *, or Unicode bullet •)
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("\u2022"):
                content = line.lstrip("-*\u2022").strip()
                self._add_rich_text_docx(doc, content, style="List Bullet")
            elif line.startswith("  - ") or line.startswith("  * "):
                content = line[4:]
                self._add_rich_text_docx(doc, content, style="List Bullet")
            # Detect project lines: "ProjectName: description..."
            elif ":" in line and not line.startswith(" "):
                colon_pos = line.index(":")
                potential_title = line[:colon_pos].strip()
                if len(potential_title) < 60 and potential_title and not potential_title[0].isdigit():
                    rest = line[colon_pos + 1:].strip()
                    para = doc.add_paragraph(style="List Bullet")
                    # Bold title
                    title_run = para.add_run(f"{potential_title}: ")
                    title_run.bold = True
                    title_run.font.size = Pt(10.5)
                    title_run.font.name = "Calibri"
                    # Rest with rich text
                    segments = self._parse_bold_segments(self._sanitize_docx_text(rest))
                    for seg_text, is_bold in segments:
                        if not seg_text:
                            continue
                        run = para.add_run(seg_text)
                        run.bold = is_bold
                        run.font.size = Pt(10.5)
                        run.font.name = "Calibri"
                else:
                    self._add_rich_text_docx(doc, line)
            else:
                self._add_rich_text_docx(doc, line)

    def _add_list_content(self, doc, items, section_name):
        """
        Add list content to the DOCX document. Handles both simple string
        lists and lists of dictionaries (like experience or education entries).

        Args:
            doc: python-docx Document object
            items (list): List of strings or dicts
            section_name (str): Section name for formatting context
        """
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for item in items:
            if isinstance(item, dict):
                # Format based on section type
                if section_name == "experience":
                    self._format_experience_docx(doc, item)
                elif section_name == "education":
                    self._format_education_docx(doc, item)
                elif section_name in ("certifications", "projects"):
                    self._format_generic_dict_docx(doc, item)
                else:
                    self._format_generic_dict_docx(doc, item)
            elif isinstance(item, str) and item and not self._is_placeholder(item):
                para = doc.add_paragraph(item, style="List Bullet")
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Calibri"

    def _add_dict_content(self, doc, data):
        """
        Add dictionary content to the DOCX (used for skills section).
        Category labels are bold, values are regular weight.

        Args:
            doc: python-docx Document object
            data (dict): Dictionary with category keys and list values
        """
        from docx.shared import Pt

        for category, values in data.items():
            if not values:
                continue
            category_label = category.replace("_", " ").title()
            if isinstance(values, list):
                values_text = ", ".join(str(v) for v in values)
            else:
                values_text = str(values)
            para = doc.add_paragraph()
            label_run = para.add_run(f"{category_label}: ")
            label_run.bold = True
            label_run.font.size = Pt(10.5)
            label_run.font.name = "Calibri"
            value_run = para.add_run(values_text)
            value_run.font.size = Pt(10.5)
            value_run.font.name = "Calibri"

    def _format_experience_docx(self, doc, exp):
        """
        Format a single work experience entry in the DOCX.

        Args:
            doc: python-docx Document object
            exp (dict): Experience dict with title, company, dates, etc.
        """
        from docx.shared import Pt, RGBColor

        title = exp.get("title", "")
        company = exp.get("company", "")
        start = exp.get("start_date", "")
        end = exp.get("end_date", "")
        dates = f"{start} - {end}" if start else ""

        # Title line: "Job Title | Company    Dates"
        header_para = doc.add_paragraph()
        header_para.space_before = Pt(4)
        header_para.space_after = Pt(1)
        title_run = header_para.add_run(f"{title} | {company}")
        title_run.bold = True
        title_run.font.size = Pt(10.5)
        title_run.font.name = "Calibri"
        if dates:
            dates_run = header_para.add_run(f"    {dates}")
            dates_run.font.size = Pt(9)
            dates_run.font.name = "Calibri"
            dates_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Description
        desc = exp.get("description", "")
        if desc:
            para = doc.add_paragraph(desc)
            for run in para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"

        # Achievement bullets
        achievements = exp.get("achievements", [])
        for ach in achievements:
            if ach:
                para = doc.add_paragraph(ach, style="List Bullet")
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Calibri"

    def _format_education_docx(self, doc, edu):
        """
        Format a single education entry in the DOCX.

        Args:
            doc: python-docx Document object
            edu (dict): Education dict with degree, institution, year, etc.
        """
        from docx.shared import Pt, RGBColor

        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year = edu.get("year", "")

        # Skip entries with no meaningful degree (e.g., certifications that leaked into education)
        if self._is_placeholder(degree) or not degree.strip():
            return

        header_para = doc.add_paragraph()
        header_para.space_before = Pt(4)
        header_para.space_after = Pt(1)
        deg_run = header_para.add_run(f"{degree} | {institution}")
        deg_run.bold = True
        deg_run.font.size = Pt(10.5)
        deg_run.font.name = "Calibri"
        if year:
            year_run = header_para.add_run(f"    {year}")
            year_run.font.size = Pt(9)
            year_run.font.name = "Calibri"
            year_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        gpa = edu.get("gpa", "")
        if gpa:
            from utils import format_gpa_label
            gpa_display = format_gpa_label(gpa)
            if gpa_display:
                para = doc.add_paragraph(gpa_display)
                for run in para.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = "Calibri"

        details = edu.get("details", "")
        if details:
            para = doc.add_paragraph(details)
            for run in para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"

    def _format_generic_dict_docx(self, doc, item):
        """
        Format a generic dictionary item as a bullet point in DOCX.

        Args:
            doc: python-docx Document object
            item (dict): Dictionary with arbitrary keys
        """
        from docx.shared import Pt

        # Handle achievement-style dicts with "title" key
        title = item.get("title", "")
        if title and not self._is_placeholder(title):
            para = doc.add_paragraph(title, style="List Bullet")
            for run in para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"
            return

        # Build a single line from the dict
        name = item.get("name", "")
        desc = item.get("description", "")
        issuer = item.get("issuer", "")
        year = item.get("year", "")
        techs = item.get("technologies", [])

        parts = []
        if name:
            parts.append(name)
        if issuer:
            parts.append(issuer)
        if year:
            parts.append(f"({year})")
        if desc:
            parts.append(f"- {desc}")
        if techs and isinstance(techs, list):
            parts.append(f"[{', '.join(techs)}]")

        text = " ".join(parts) if parts else json.dumps(item)
        para = doc.add_paragraph(text, style="List Bullet")
        for run in para.runs:
            run.font.size = Pt(10.5)
            run.font.name = "Calibri"

    # =========================================================================
    # PDF Helper Methods
    # =========================================================================

    @staticmethod
    def _sanitize_pdf_text(text):
        """
        Make text safe for fpdf2's built-in Helvetica font, which only supports
        the Latin-1 (ISO-8859-1) character set.

        Replacements performed:
        - Bullet characters (•, ▪, ◦, ‣, ➤, ➢) → ASCII "- "
        - Markdown italic (*text*) and underline (__text__) → plain text
        - Markdown bold (**text**) → PRESERVED for rich text rendering
        - En-dash / em-dash → hyphen
        - Smart quotes → straight quotes
        - Ellipsis character → three dots
        - Any remaining non-Latin-1 characters → stripped out

        Args:
            text (str): Raw text that may contain Unicode characters

        Returns:
            str: Latin-1-safe text for PDF rendering
        """
        if not text:
            return text

        import re

        # Replace bullet-like characters at the start of a word with "- "
        text = re.sub(r'[•▪◦‣➤➢]\s*', '- ', text)

        # Strip markdown italic markers (*text* or __text__) but KEEP **bold**
        # Bold markers are preserved for rich text rendering in _add_rich_text_pdf
        text = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)

        # Normalize dashes and quotes
        text = text.replace('\u2013', '-')   # en-dash
        text = text.replace('\u2014', '-')   # em-dash
        text = text.replace('\u2018', "'")   # left single quote
        text = text.replace('\u2019', "'")   # right single quote / apostrophe
        text = text.replace('\u201c', '"')   # left double quote
        text = text.replace('\u201d', '"')   # right double quote
        text = text.replace('\u2026', '...')  # ellipsis

        # Strip any remaining characters outside Latin-1 range
        text = text.encode('latin-1', errors='ignore').decode('latin-1')

        return text

    def _add_rich_text_pdf(self, pdf, text, indent=""):
        """
        Render text with **bold** markers as actual bold segments in the PDF.
        Uses fpdf2's write() method with font toggling to mix bold and regular
        text on the same line.  Falls back to plain multi_cell for text
        that has no bold markers (avoids unnecessary overhead).

        Args:
            pdf: FPDF object
            text (str): Text potentially containing **bold** markers
            indent (str): Leading characters (e.g. "  - ") prepended before the text
        """
        # Sanitize the text first (preserves **bold** markers)
        sanitized = self._sanitize_pdf_text(text)
        segments = self._parse_bold_segments(sanitized)

        # Fast path: no bold markers — use simple multi_cell
        has_bold = any(is_bold for _, is_bold in segments)
        if not has_bold:
            full_text = indent + sanitized if indent else sanitized
            pdf.multi_cell(0, 5, full_text, new_x="LMARGIN", new_y="NEXT")
            return

        # Write the indent first (regular font)
        if indent:
            pdf.set_font("Helvetica", "", 10)
            pdf.write(5, indent)

        # Write each segment, toggling bold on/off
        for segment_text, is_bold in segments:
            if not segment_text:
                continue
            style = "B" if is_bold else ""
            pdf.set_font("Helvetica", style, 10)
            pdf.write(5, segment_text)

        # Move to the next line after the mixed-format line
        pdf.set_font("Helvetica", "", 10)
        pdf.ln(5)

    def _add_text_pdf(self, pdf, text, section_name=""):
        """
        Add text content to the PDF. Handles multi-line text and bullet points.
        Sanitizes each line to Latin-1 before rendering.
        Detects experience title lines and renders them fully bold.
        Renders **bold** markers as actual bold text in bullet points
        (for keywords like technologies, tools, and metrics).

        Args:
            pdf: FPDF object
            text (str): Text content to add
            section_name (str): Section name for context-aware formatting
        """
        import re
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Skip placeholder lines (NONE, N/A, etc.)
            if self._is_placeholder(line):
                continue
            # Skip markdown horizontal rules (---, ***, ___)
            if set(line) <= {'-', '*', '_'} and len(line) >= 3:
                continue
            # Detect experience title lines (e.g. "Title | Company | Date")
            # and render them bold with a small top margin
            if self._is_experience_title_line(line):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 10)
                pdf.multi_cell(0, 5, self._sanitize_pdf_text(line),
                               new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", "", 10)
            # Detect bullet lines: -, *, or bullet char at start
            elif line.startswith("- ") or line.startswith("* ") or line.startswith("\u2022"):
                content = line.lstrip("-*\u2022").strip()
                self._add_rich_text_pdf(pdf, content, indent="  - ")
            # Detect project lines: "ProjectName: description..."
            elif ":" in line and not line.startswith(" "):
                colon_pos = line.index(":")
                potential_title = line[:colon_pos].strip()
                # Project titles are typically short (under 60 chars) and don't start with bullets
                if len(potential_title) < 60 and potential_title and not potential_title[0].isdigit():
                    rest = line[colon_pos + 1:].strip()
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.write(5, f"  - {potential_title}: ")
                    pdf.set_font("Helvetica", "", 10)
                    # Render the rest with rich text for **bold** keywords
                    segments = self._parse_bold_segments(self._sanitize_pdf_text(rest))
                    for seg_text, is_bold in segments:
                        if not seg_text:
                            continue
                        style = "B" if is_bold else ""
                        pdf.set_font("Helvetica", style, 10)
                        pdf.write(5, seg_text)
                    pdf.set_font("Helvetica", "", 10)
                    pdf.ln(5)
                else:
                    self._add_rich_text_pdf(pdf, line)
            else:
                # Regular line — also supports bold keywords
                self._add_rich_text_pdf(pdf, line)

    def _add_list_pdf(self, pdf, items, section_name):
        """
        Add list content to the PDF.

        Args:
            pdf: FPDF object
            items (list): List of strings or dicts
            section_name (str): Section name for formatting context
        """
        for item in items:
            if isinstance(item, dict):
                if section_name == "experience":
                    self._format_experience_pdf(pdf, item)
                elif section_name == "education":
                    self._format_education_pdf(pdf, item)
                else:
                    self._format_generic_dict_pdf(pdf, item)
            elif isinstance(item, str) and item and not self._is_placeholder(item):
                pdf.multi_cell(0, 5, self._sanitize_pdf_text(f"  - {item}"),
                               new_x="LMARGIN", new_y="NEXT")

    def _add_dict_pdf(self, pdf, data):
        """
        Add dictionary content to the PDF (used for skills section).
        Category labels are bold, values are regular weight.

        Args:
            pdf: FPDF object
            data (dict): Dictionary with category keys and list values
        """
        for category, values in data.items():
            if not values:
                continue
            category_label = category.replace("_", " ").title()
            if isinstance(values, list):
                values_text = ", ".join(str(v) for v in values)
            else:
                values_text = str(values)
            pdf.set_font("Helvetica", "B", 10)
            pdf.write(5, self._sanitize_pdf_text(f"{category_label}: "))
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, self._sanitize_pdf_text(values_text),
                           new_x="LMARGIN", new_y="NEXT")

    def _format_experience_pdf(self, pdf, exp):
        """
        Format a single work experience entry in the PDF.

        Args:
            pdf: FPDF object
            exp (dict): Experience dict
        """
        title = exp.get("title", "")
        company = exp.get("company", "")
        start = exp.get("start_date", "")
        end = exp.get("end_date", "")
        dates = f"{start} - {end}" if start else ""

        # Title line
        pdf.set_font("Helvetica", "B", 10)
        header = f"{title} | {company}"
        if dates:
            header += f"    {dates}"
        pdf.multi_cell(0, 5, self._sanitize_pdf_text(header),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)

        # Description
        desc = exp.get("description", "")
        if desc:
            pdf.multi_cell(0, 5, self._sanitize_pdf_text(desc),
                           new_x="LMARGIN", new_y="NEXT")

        # Achievements
        achievements = exp.get("achievements", [])
        for ach in achievements:
            if ach:
                pdf.multi_cell(0, 5, self._sanitize_pdf_text(f"  - {ach}"),
                               new_x="LMARGIN", new_y="NEXT")

        pdf.ln(2)

    def _format_education_pdf(self, pdf, edu):
        """
        Format a single education entry in the PDF.

        Args:
            pdf: FPDF object
            edu (dict): Education dict
        """
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year = edu.get("year", "")

        # Skip entries with no meaningful degree (e.g., certifications that leaked into education)
        if self._is_placeholder(degree) or not degree.strip():
            return

        pdf.set_font("Helvetica", "B", 10)
        header = f"{degree} | {institution}"
        if year:
            header += f"    {year}"
        pdf.multi_cell(0, 5, self._sanitize_pdf_text(header),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)

        gpa = edu.get("gpa", "")
        if gpa:
            from utils import format_gpa_label
            gpa_display = format_gpa_label(gpa)
            if gpa_display:
                pdf.multi_cell(0, 5, self._sanitize_pdf_text(gpa_display),
                               new_x="LMARGIN", new_y="NEXT")

        pdf.ln(1)

    def _format_generic_dict_pdf(self, pdf, item):
        """
        Format a generic dictionary item as a bullet point in PDF.

        Args:
            pdf: FPDF object
            item (dict): Dictionary with arbitrary keys
        """
        # Handle achievement-style dicts with "title" key
        title = item.get("title", "")
        if title and not self._is_placeholder(title):
            pdf.multi_cell(0, 5, self._sanitize_pdf_text(f"  - {title}"),
                           new_x="LMARGIN", new_y="NEXT")
            return

        name = item.get("name", "")
        desc = item.get("description", "")
        issuer = item.get("issuer", "")
        year = item.get("year", "")
        techs = item.get("technologies", [])

        parts = []
        if name:
            parts.append(name)
        if issuer:
            parts.append(issuer)
        if year:
            parts.append(f"({year})")
        if desc:
            parts.append(f"- {desc}")
        if techs and isinstance(techs, list):
            parts.append(f"[{', '.join(techs)}]")

        text = " ".join(parts) if parts else json.dumps(item)
        pdf.multi_cell(0, 5, self._sanitize_pdf_text(f"  - {text}"),
                       new_x="LMARGIN", new_y="NEXT")
